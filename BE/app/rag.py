from __future__ import annotations

import os
import re
import subprocess
import uuid
from collections import defaultdict, deque
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from dotenv import load_dotenv
import httpx
from openai import OpenAI
from pypdf import PdfReader

# Resolve configuration relative to the backend, not the process working
# directory. This works equally with `uvicorn app.main:app` from BE and
# `uvicorn BE.app.main:app` from the project root.
BACKEND_DIR = Path(__file__).resolve().parents[1]
ENV_FILE = BACKEND_DIR / ".env"
ENV_EXAMPLE_FILE = BACKEND_DIR / ".env.example"
load_dotenv(dotenv_path=ENV_FILE if ENV_FILE.exists() else ENV_EXAMPLE_FILE, override=False)

# config.py and the checked-in cache use this folder name (without a leading
# dot). It is a tokenizer cache, not an SSL certificate or embedding model.
TIKTOKEN_CACHE_DIR = BACKEND_DIR / "tiktoken_cache"
if TIKTOKEN_CACHE_DIR.is_dir():
    os.environ.setdefault("TIKTOKEN_CACHE_DIR", str(TIKTOKEN_CACHE_DIR))


def configure_ca_bundle() -> None:
    """Make a corporate/self-signed CA available to Hugging Face HTTPS clients."""
    certificate = os.getenv("SSL_CERT_FILE") or os.getenv("REQUESTS_CA_BUNDLE")
    if not certificate:
        return
    path = Path(certificate)
    if not path.is_absolute():
        path = BACKEND_DIR / path
    if not path.is_file():
        raise ValueError(f"SSL certificate bundle was not found: {path}")
    resolved = str(path.resolve())
    # requests, curl, httpx, and libraries layered on them use one of these.
    os.environ["SSL_CERT_FILE"] = resolved
    os.environ["REQUESTS_CA_BUNDLE"] = resolved
    os.environ["CURL_CA_BUNDLE"] = resolved


configure_ca_bundle()


class RAGPipeline:
    """Loader -> chunks -> embeddings -> Chroma -> retrieve -> limit -> LLM -> memory."""

    def __init__(self) -> None:
        self.upload_path = Path(os.getenv("UPLOAD_PATH", "./data/uploads"))
        self.upload_path.mkdir(parents=True, exist_ok=True)
        # Support both this service's OpenAI names and the GenAI Lab names
        # already used by config.py.
        key = os.getenv("OPENAI_API_KEY") or os.getenv("GENAILAB_API_KEY")
        base_url = os.getenv("OPENAI_BASE_URL") or os.getenv("GENAILAB_BASE_URL")
        verify_ssl = os.getenv("OPENAI_VERIFY_SSL", "true").strip().lower() not in {"0", "false", "no"}
        self.http_client = httpx.Client(verify=verify_ssl) if not verify_ssl else None
        self.llm = OpenAI(api_key=key, base_url=base_url or None, http_client=self.http_client) if key else None
        self.llm_model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        self.embedding_provider = os.getenv("EMBEDDING_PROVIDER", "sentence_transformers").lower()
        self.embedding_model = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
        if self.embedding_provider == "openai":
            if not self.llm:
                raise ValueError("OPENAI_API_KEY is required when EMBEDDING_PROVIDER=openai.")
            self.model = None
        elif self.embedding_provider == "sentence_transformers":
            # Import only for the local-model option. This avoids loading
            # PyTorch/Hugging Face during OpenAI-compatible API startup.
            from sentence_transformers import SentenceTransformer

            self.model = SentenceTransformer(self.embedding_model)
        else:
            raise ValueError("EMBEDDING_PROVIDER must be 'openai' or 'sentence_transformers'.")
        # Delay Chroma's heavier import until the first RAG request so the
        # FastAPI health/docs endpoints can start immediately.
        import chromadb

        client = chromadb.PersistentClient(path=os.getenv("CHROMA_PATH", "./data/chroma"))
        self.collection = client.get_or_create_collection(
            name="documents", metadata={"hnsw:space": "cosine"}
        )
        self.memory: dict[str, deque[dict[str, str]]] = defaultdict(lambda: deque(maxlen=8))


    def embed(self, texts: list[str]) -> list[list[float]]:
        """Create vectors through a local Hugging Face model or an OpenAI-compatible API."""
        if self.embedding_provider == "openai":
            if not self.llm:  # guarded during initialization; keeps type checkers happy.
                raise ValueError("OpenAI embedding client is not configured.")
            response = self.llm.embeddings.create(model=self.embedding_model, input=texts)
            return [item.embedding for item in response.data]
        if not self.model:
            raise ValueError("SentenceTransformer embedding model is not configured.")
        return self.model.encode(texts, normalize_embeddings=True).tolist()

    @staticmethod
    def _clean(text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def load(self, path: Path) -> list[tuple[int, str]]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return [(index + 1, self._clean(page.extract_text() or "")) for index, page in enumerate(reader.pages)]
        if suffix == ".docx":
            document = DocxDocument(str(path))
            return [(1, self._clean("\n".join(p.text for p in document.paragraphs)))]
        if suffix == ".doc":
            # Legacy Word binary files need a converter; LibreOffice is a widely available option.
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(path.parent), str(path)],
                    check=True, capture_output=True, timeout=60,
                )
            except (FileNotFoundError, subprocess.SubprocessError) as error:
                raise ValueError("Legacy .doc files require LibreOffice (soffice) on the server; convert to .docx or install LibreOffice.") from error
            converted = path.with_suffix(".docx")
            if not converted.exists():
                raise ValueError("LibreOffice could not convert this .doc file.")
            try:
                document = DocxDocument(str(converted))
                return [(1, self._clean("\n".join(p.text for p in document.paragraphs)))]
            finally:
                converted.unlink(missing_ok=True)
        if suffix == ".txt":
            return [(1, self._clean(path.read_text(encoding="utf-8", errors="ignore")))]
        raise ValueError("Only PDF, DOC, DOCX, and TXT files are supported.")

    @staticmethod
    def chunk(text: str, size: int = 900, overlap: int = 160) -> Iterable[str]:
        if not text:
            return
        start = 0
        while start < len(text):
            end = min(start + size, len(text))
            if end < len(text):
                split = max(text.rfind(". ", start, end), text.rfind(" ", start, end))
                if split > start + size // 2:
                    end = split + 1
            yield text[start:end].strip()
            if end == len(text):
                break
            start = end - overlap

    def ingest(self, filename: str, content: bytes) -> dict[str, object]:
        safe_name = Path(filename).name
        document_id = uuid.uuid4().hex
        saved = self.upload_path / f"{document_id}_{safe_name}"
        saved.write_bytes(content)
        try:
            pages = self.load(saved)
        except Exception:
            saved.unlink(missing_ok=True)
            raise

        ids, texts, metadata = [], [], []
        for page_number, page_text in pages:
            for index, text in enumerate(self.chunk(page_text)):
                if text:
                    ids.append(f"{document_id}-{page_number}-{index}")
                    texts.append(text)
                    metadata.append({"document_id": document_id, "filename": safe_name, "page": page_number})
        if not texts:
            raise ValueError("No readable text was found in this document.")
        embeddings = self.embed(texts)
        self.collection.add(ids=ids, documents=texts, metadatas=metadata, embeddings=embeddings)
        return {"document_id": document_id, "filename": safe_name, "chunks": len(texts)}

    def retrieve(self, question: str, document_ids: list[str] | None, limit: int = 5) -> list[dict[str, object]]:
        where = {"document_id": {"$in": document_ids}} if document_ids else None
        embedding = self.embed([question])
        result = self.collection.query(query_embeddings=embedding, n_results=limit, where=where, include=["documents", "metadatas", "distances"])
        return [
            {"text": text, "source": metadata["filename"], "page": metadata["page"], "score": round(1 - distance, 3)}
            for text, metadata, distance in zip(result["documents"][0], result["metadatas"][0], result["distances"][0])
        ]

    @staticmethod
    def limit_context(hits: list[dict[str, object]], characters: int = 6000) -> str:
        selected, used = [], 0
        for hit in hits:
            excerpt = str(hit["text"])
            remaining = characters - used
            if remaining <= 0:
                break
            selected.append(f"[{hit['source']} p.{hit['page']}] {excerpt[:remaining]}")
            used += len(excerpt)
        return "\n\n".join(selected)

    def ask(self, question: str, session_id: str, document_ids: list[str] | None = None) -> dict[str, object]:
        hits = self.retrieve(question, document_ids)
        context = self.limit_context(hits)
        history = "\n".join(f"{m['role']}: {m['content']}" for m in self.memory[session_id])
        if self.llm:
            prompt = f"""Answer only from the supplied document context. If it is insufficient, say so plainly.
Conversation memory:\n{history}\n\nDocument context:\n{context}\n\nQuestion: {question}"""
            response = self.llm.chat.completions.create(
                model=self.llm_model,
                messages=[{"role": "system", "content": "You are a precise RAG assistant. Cite filename and page."}, {"role": "user", "content": prompt}],
                temperature=0.2,
            )
            answer = response.choices[0].message.content or "I could not generate an answer."
        elif hits:
            answer = "LLM is not configured. The most relevant document passages are returned below; set OPENAI_API_KEY to enable synthesized answers."
        else:
            answer = "I could not find relevant information in the indexed documents."
        self.memory[session_id].append({"role": "user", "content": question})
        self.memory[session_id].append({"role": "assistant", "content": answer})
        return {"answer": answer, "sources": hits, "context_characters": len(context)}
